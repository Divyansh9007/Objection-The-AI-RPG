from docx import Document
from typing import Dict, List
import json
import os

class OutputGenerator:
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def save_json(self, data: Dict, filename: str):
        """Save results as JSON file."""
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
            
    def create_word_report(self, summary_data: Dict, filename: str):
        """Generate a Word document with formatted summaries."""
        doc = Document()
        
        # Add executive summary
        doc.add_heading('Contract Summary Report', 0)
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(summary_data["executive_summary"])
        
        # Add important clauses section
        doc.add_heading('Important Clauses', level=1)
        for clause in summary_data["clause_summaries"]:
            doc.add_heading(f'{clause["type"].title()} Clause', level=2)
            doc.add_paragraph(f'Heading: {clause["heading"]}')
            doc.add_paragraph(f'Summary: {clause["summary"]}')
            doc.add_paragraph(f'Confidence: {clause["confidence"]:.2%}')
            
        # Save document
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
